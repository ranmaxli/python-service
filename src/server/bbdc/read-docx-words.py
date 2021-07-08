#!/usr/bin/python
# -*- coding: UTF-8 -*-

# 名称：该软件主要将我们日常所需要的记忆的词汇转换成“不背单词”自定义词汇所需要的文件格式
# 时间：2021-7-9 07:33:46
# 作者：ranmaxli

# # 使用说明
# # 1、找到一篇喜欢的文章，把文章内容复制下来，粘贴到docx文件下，并放置项目word-dictionary文件夹下。
# # 2、更改读取的 get_word 方法中的 file=docx.Document(" URL ") 的地址
# # 3、运行程序后可以到word-txt取名称为最新的txt文本上传至 bbdc.cn 取得最新词汇表
# #
# # # 项目结构说明
# # # word-dictionary
# #     用于存放用户自己整理的word单词文档
# # # word-txt
# #     系统分析后生成的单词字典表
# # # everyday-vocabulary.txt
# #     常用的词汇，如果你想屏蔽一些已经会的词汇，直接添加在这个文件中即可，系统分析时会直接跳过
# # # read-docx-words.py
# #     项目主函数入口

import docx
import re
from collections import Counter
import numpy as np
import datetime
import pytz
import os

def get_word():
    word_list = []
    file=docx.Document(os.getcwd() + "/word-dictionary/cet4-composition.docx")
    for i in range(len(file.paragraphs)):
        result_list = re.findall("[a-zA-Z]+",file.paragraphs[i].text)
        for word in result_list:
            word_list.append(word)
    return word_list

def remove_everyday_vocabulary(word_list_desc):
    file = open(os.getcwd() + '/everyday-vocabulary.txt')
    everyday_vocabulary = file.readlines()
    everyday_vocabulary = [' '.join([i.strip() for i in price.strip().split('\n')]) for price in everyday_vocabulary] #去除换行符
    everyday_vocabulary = Counter(everyday_vocabulary) #去除相同元素
    everyday_vocabulary = [i for i in everyday_vocabulary if i != ''] #去除空值
    everyday_vocabulary_capitalization = [s.capitalize() for s in everyday_vocabulary if isinstance(s,str)==True] #生成首字母大写单词
    everyday_vocabulary.extend(everyday_vocabulary_capitalization) 

    for word in everyday_vocabulary:
        if word in word_list_desc:
            word_list_desc.remove(word)
    return word_list_desc

def analysis_word(word_list):
    word_list_Counter = Counter(word_list)
    word_list_desc_tuple = sorted(word_list_Counter.items(), key=lambda x: x[1], reverse=True)
    word_list_desc = []
    for word in word_list_desc_tuple:
        word_list_desc.append(word[0])
        # Fix translate word use ThreadPoolExecutor
    word_list_desc_distinct = remove_everyday_vocabulary(word_list_desc)
    file_name = datetime.datetime.now(pytz.timezone('PRC')).strftime('%Y-%m-%d-%H-%M-%S')
    for word in word_list_desc_distinct:
        write_word_to_file(word,file_name)
        print(word)

    print("共计          词汇数量  ：" + str(len(word_list)))
    print("去重后        词汇数量  ：" + str(len(list(set(word_list)))))
    print("去除常用词汇后 词汇数量  ：" + str(len(word_list_desc_distinct)))
    print("按照词汇出现频率从高到低排序")

def write_word_to_file(word,file_name):
    txt_file_directory_address = os.getcwd() + '/word-txt/' + file_name +'.txt'
    f=open(txt_file_directory_address, "a+")
    f.write(word + "\n")   
    f.close()

if __name__ == "__main__":
    analysis_word(get_word())